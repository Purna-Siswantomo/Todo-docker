# -*- coding: utf-8 -*-
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent
REPO = BASE.parent
EVIDENCE = REPO / "evidence"
DIAGRAM = BASE / "architecture-diagram.png"

ANSI_RE = re.compile(r"\x1b\[[0-9;]*m")


def strip_ansi(text: str) -> str:
    return ANSI_RE.sub("", text)


def read_evidence(name: str) -> str:
    return strip_ansi((EVIDENCE / name).read_text(encoding="utf-8"))


doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for level, size, color in [(1, 16, "1F4E79"), (2, 13, "1F4E79"), (3, 12, "1F4E79")]:
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    st.paragraph_format.space_after = Pt(6)

sections = doc.sections
for s in sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2.5)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


footer = sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number_field(fp)


def code_block(text: str, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    border = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "BFBFBF")
        border.append(el)
    pPr.append(border)
    lines = text.rstrip("\n").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(size)
        if i != len(lines) - 1:
            run.add_break()
    return p


def heading(text, level=1):
    doc.add_heading(text, level=level)


def para(text, bold=False, italic=False, align=None, size=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def page_break():
    doc.add_page_break()


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = para("LAPORAN UJIAN AKHIR SEMESTER", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
p = para("PROYEK AKHIR CLOUD COMPUTING", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
doc.add_paragraph()
p = para(
    "IMPLEMENTASI APLIKASI MULTI-CONTAINER DENGAN DOCKER, ORKESTRASI, DAN CI/CD",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18,
)
doc.add_paragraph()
p = para("Studi Kasus: Todo List Application (Laravel + MariaDB)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)

for _ in range(4):
    doc.add_paragraph()

identity_lines = [
    "Disusun oleh:",
    "Purna Siswantomo  -  C2C023160",
    "Erifa Dwi Astuti   -  C2C023161",
    "Kelas: [Kelas Anda]",
    "",
    "Mata Kuliah: Komputasi Awan (Cloud Computing)",
    "Tahun Akademik: 2025/2026",
]
for line in identity_lines:
    para(line, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=(line == "Disusun oleh:"))

for _ in range(3):
    doc.add_paragraph()
para("Repository: https://github.com/Purna-Siswantomo/Todo-docker", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

page_break()

# ============================================================
# DAFTAR ISI (manual, tanpa nomor halaman otomatis)
# ============================================================
heading("DAFTAR ISI", level=1)
toc_entries = [
    "BAB 1  PENDAHULUAN",
    "        1.1 Latar Belakang",
    "        1.2 Permasalahan",
    "        1.3 Tujuan",
    "        1.4 Manfaat",
    "BAB 2  ANALISIS DAN ARSITEKTUR",
    "        2.1 Deskripsi Aplikasi",
    "        2.2 Pengguna Sistem",
    "        2.3 Teknologi yang Digunakan",
    "        2.4 Diagram Arsitektur",
    "        2.5 Fungsi Setiap Komponen",
    "BAB 3  IMPLEMENTASI APLIKASI",
    "        3.1 Fitur Utama",
    "        3.2 Struktur Data",
    "        3.3 Implementasi Model dan Controller",
    "        3.4 Routing dan Endpoint",
    "        3.5 Konfigurasi Environment",
    "BAB 4  IMPLEMENTASI CONTAINER",
    "        4.1 Dockerfile",
    "        4.2 Docker Compose",
    "        4.3 Network dan Komunikasi Antar Service",
    "        4.4 Persistent Volume",
    "        4.5 Health Check dan Ketahanan Layanan",
    "BAB 5  IMPLEMENTASI CI/CD",
    "        5.1 Tujuan Pipeline",
    "        5.2 Alur Workflow",
    "        5.3 Automated Testing sebagai Quality Gate",
    "        5.4 Simulasi Pipeline Gagal dan Berhasil",
    "BAB 6  PENGUJIAN",
    "        6.1 Pengujian Fungsional",
    "        6.2 Pengujian Automated Test",
    "        6.3 Pengujian Container dan Compose",
    "        6.4 Pengujian Persistent Volume",
    "        6.5 Pengujian Ketahanan Layanan (Health Check)",
    "BAB 7  KESIMPULAN",
    "        7.1 Kesimpulan Umum",
    "        7.2 Kendala dan Solusi",
    "        7.3 Rencana Pengembangan",
    "LAMPIRAN",
]
for entry in toc_entries:
    p = para(entry, size=10)
    p.paragraph_format.space_after = Pt(1)

# ============================================================
# BAB 1 PENDAHULUAN
# ============================================================
heading("BAB 1 PENDAHULUAN", level=1)

heading("1.1 Latar Belakang", level=2)
para(
    "Perkembangan cloud computing mendorong kebutuhan akan aplikasi yang mudah dipindahkan, "
    "konsisten antar lingkungan, dan mudah dikelola. Salah satu pendekatan yang banyak digunakan "
    "untuk menjawab kebutuhan tersebut adalah containerization menggunakan Docker. Dengan Docker, "
    "sebuah aplikasi dapat dikemas bersama dependensinya ke dalam image yang sama, sehingga perilaku "
    "aplikasi menjadi lebih stabil ketika dijalankan pada mesin yang berbeda. Pada proyek ini, konsep "
    "tersebut diterapkan melalui pengembangan aplikasi Todo List berbasis Laravel yang dijalankan di "
    "dalam arsitektur multi-container."
)

heading("1.2 Permasalahan", level=2)
para(
    "Sebelum menggunakan container, aplikasi sering mengalami perbedaan environment antara mesin "
    "pengembang, penguji, dan production. Perbedaan versi dependency, konfigurasi database, serta "
    "proses instalasi manual dapat menimbulkan error yang sulit dilacak. Selain itu, jika database "
    "dijalankan tanpa persistent volume, data berisiko hilang saat container dihentikan. Permasalahan "
    "lain adalah kurangnya otomasi dalam pengujian dan build, sehingga perubahan kode tidak selalu "
    "diverifikasi sebelum diterapkan. Hal-hal tersebut menjadi alasan mengapa pendekatan Docker, Docker "
    "Compose, volume, health check, dan CI/CD perlu diterapkan dalam proyek ini."
)

heading("1.3 Tujuan", level=2)
para(
    "Tujuan proyek ini adalah membangun aplikasi sederhana yang berjalan dalam arsitektur "
    "multi-container dengan container aplikasi dan container database. Proyek ini juga bertujuan untuk "
    "menerapkan persistent volume agar data tetap aman, menggunakan environment variable untuk "
    "konfigurasi yang fleksibel, menambahkan health check untuk pemantauan layanan, serta menerapkan "
    "automated testing dan GitHub Actions sebagai pipeline CI/CD. Dengan demikian, proyek ini tidak "
    "hanya menghasilkan aplikasi fungsional, tetapi juga memperlihatkan alur kerja deployment modern "
    "yang sesuai dengan konsep cloud computing."
)

heading("1.4 Manfaat", level=2)
para(
    "Manfaat dari proyek ini adalah memberikan pemahaman praktis mengenai hubungan antara aplikasi, "
    "image, container, orchestrator, network, dan volume. Selain itu, proyek ini menunjukkan bagaimana "
    "otomasi testing dan build membantu menjaga kualitas aplikasi. Dari sisi praktis, penggunaan Docker "
    "membuat aplikasi lebih mudah dijalankan ulang, sementara volume menjaga data database tetap "
    "tersedia. Hasil akhir proyek ini juga dapat menjadi dasar pengembangan ke arah deployment yang "
    "lebih profesional pada layanan cloud atau VPS."
)


# ============================================================
# BAB 2 ANALISIS DAN ARSITEKTUR
# ============================================================
heading("BAB 2 ANALISIS DAN ARSITEKTUR", level=1)

heading("2.1 Deskripsi Aplikasi", level=2)
para(
    "Aplikasi yang dikembangkan pada proyek ini adalah Todo List sederhana yang digunakan untuk "
    "mengelola data tugas. Pengguna dapat melihat daftar todo, membuat todo baru, melihat detail todo, "
    "memperbarui todo, dan menghapus todo yang sudah tidak diperlukan. Walaupun fiturnya sederhana, "
    "aplikasi ini dirancang untuk memperlihatkan implementasi konsep cloud computing secara utuh melalui "
    "penggunaan container, database terpisah, health check, dan pipeline otomatis."
)

heading("2.2 Pengguna Sistem", level=2)
para(
    "Pengguna sistem adalah mahasiswa atau pengguna umum yang ingin mencatat dan mengelola daftar "
    "tugas. Dalam konteks pengujian, pengguna mengakses aplikasi melalui browser atau HTTP client "
    "seperti curl untuk memanggil endpoint API. Sistem kemudian memproses request tersebut melalui "
    "container aplikasi dan menyimpan atau membaca data pada database yang berjalan di container "
    "terpisah."
)

heading("2.3 Teknologi yang Digunakan", level=2)
para(
    "Teknologi utama yang digunakan adalah PHP 8.3, Laravel 13, Docker, Docker Compose, MariaDB "
    "11.4, PHPUnit, dan GitHub Actions. Laravel digunakan untuk membangun aplikasi web dan REST API. "
    "Docker digunakan untuk containerization, sedangkan Docker Compose digunakan untuk mengorkestrasi "
    "beberapa service dalam satu file konfigurasi. MariaDB dipilih sebagai database runtime, sementara "
    "SQLite digunakan untuk testing lokal yang terisolasi dari data produksi. GitHub Actions digunakan "
    "untuk menjalankan workflow otomatis berupa testing dan build image."
)

heading("2.4 Diagram Arsitektur", level=2)
para(
    "Arsitektur aplikasi terdiri dari pengguna, container aplikasi, container database, network "
    "internal, volume penyimpanan data, dan pipeline CI/CD. Pengguna mengakses aplikasi melalui port "
    "8080 pada host. Request diteruskan ke container aplikasi Laravel, lalu aplikasi berkomunikasi "
    "dengan database MariaDB melalui network Docker internal bernama todo-network. Data database "
    "disimpan pada volume db_data sehingga tetap aman saat container dihentikan atau dibuat ulang. "
    "GitHub Actions menjalankan proses checkout, instalasi dependency, automated testing, build image, "
    "dan verifikasi health check sebelum image dinyatakan layak. Diagram berikut memperlihatkan "
    "bagaimana seluruh komponen tersebut saling terhubung secara deklaratif melalui Docker Compose."
)
if DIAGRAM.exists():
    doc.add_picture(str(DIAGRAM), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Gambar 2.1 Diagram arsitektur aplikasi dan container")
    r.italic = True
    r.font.size = Pt(10)

heading("2.5 Fungsi Setiap Komponen", level=2)
para(
    "Container app berfungsi menjalankan aplikasi Laravel dan menyediakan endpoint API. Container db "
    "berfungsi menyimpan data todo secara persisten. Volume db_data menyimpan file database agar data "
    "tidak hilang saat container dihapus. Network todo-network memungkinkan container saling "
    "berkomunikasi menggunakan nama service, bukan alamat IP statis. GitHub Actions berfungsi "
    "menjalankan proses otomasi berupa testing dan build. Seluruh komponen tersebut membentuk sistem "
    "yang lebih mudah dikelola dibandingkan menjalankan semua bagian secara manual pada satu lingkungan."
)


# ============================================================
# BAB 3 IMPLEMENTASI APLIKASI
# ============================================================
heading("BAB 3 IMPLEMENTASI APLIKASI", level=1)

heading("3.1 Fitur Utama", level=2)
para(
    "Aplikasi Todo List memiliki fitur utama CRUD, yaitu create, read, update, dan delete. Pengguna "
    "dapat menambahkan data todo baru, melihat semua data todo yang tersimpan, memperbarui isi atau "
    "status todo, serta menghapus data yang tidak lagi dibutuhkan. Selain itu, aplikasi memiliki "
    "validasi input agar field penting seperti title tidak kosong. Fitur health check juga disediakan "
    "untuk memantau apakah aplikasi masih dapat terhubung ke database dengan baik."
)

heading("3.2 Struktur Data", level=2)
para(
    "Data todo disimpan dalam tabel todos. Tabel ini memiliki kolom id, title, description, "
    "completed, created_at, dan updated_at. Kolom title menyimpan nama tugas, description menyimpan "
    "keterangan tambahan, dan completed menyimpan status tugas yang telah selesai atau belum. Struktur "
    "data ini sederhana, tetapi cukup untuk membuktikan kemampuan aplikasi dalam melakukan operasi data "
    "dasar yang terhubung ke database."
)

heading("3.3 Implementasi Model dan Controller", level=2)
para(
    "Model Todo digunakan untuk mendefinisikan atribut yang dapat diisi dan mengatur cast boolean "
    "pada field completed. Controller TodoController menangani logika index, store, show, update, dan "
    "destroy. Setiap request divalidasi sebelum data disimpan atau diubah, sehingga aplikasi memiliki "
    "kontrol dasar terhadap input pengguna."
)
code_block(
    "public function store(Request $request)\n"
    "{\n"
    "    $validated = $request->validate([\n"
    "        'title' => 'required|string|max:255',\n"
    "        'description' => 'nullable|string',\n"
    "        'completed' => 'boolean',\n"
    "    ]);\n\n"
    "    $todo = Todo::create($validated);\n\n"
    "    return response()->json($todo, 201);\n"
    "}"
)

heading("3.4 Routing dan Endpoint", level=2)
para(
    "Routing aplikasi didefinisikan pada routes/api.php. Endpoint /api/todos digunakan untuk akses "
    "data todo, sedangkan endpoint /api/health digunakan untuk memeriksa kesehatan layanan dengan cara "
    "mencoba membuka koneksi PDO ke database. Route resource digunakan agar endpoint CRUD dapat dibuat "
    "secara ringkas."
)
code_block(
    "Route::get('/health', function () {\n"
    "    try {\n"
    "        DB::connection()->getPdo();\n"
    "        return response()->json(['status' => 'healthy']);\n"
    "    } catch (\\Exception $e) {\n"
    "        return response()->json(['status' => 'unhealthy', 'error' => $e->getMessage()], 503);\n"
    "    }\n"
    "});\n\n"
    "Route::apiResource('todos', TodoController::class);"
)

heading("3.5 Konfigurasi Environment", level=2)
para(
    "Konfigurasi aplikasi menggunakan file .env.example sebagai template environment. File ini berisi "
    "parameter seperti DB_CONNECTION, DB_HOST, DB_PORT, DB_DATABASE, DB_USERNAME, DB_PASSWORD, dan "
    "DB_ROOT_PASSWORD. File .env yang berisi nilai aktual tidak diunggah ke repository (didaftarkan "
    "pada .gitignore), sehingga credential tidak pernah tersimpan langsung di riwayat commit. "
    "docker-compose.yml sendiri hanya mereferensikan variabel tersebut melalui sintaks ${VAR}, bukan "
    "menuliskan nilai kredensial secara eksplisit."
)


# ============================================================
# BAB 4 IMPLEMENTASI CONTAINER
# ============================================================
heading("BAB 4 IMPLEMENTASI CONTAINER", level=1)

heading("4.1 Dockerfile", level=2)
para(
    "Dockerfile digunakan untuk membangun image aplikasi Laravel. Base image yang dipakai adalah "
    "php:8.3-cli-alpine karena ringan dan sesuai untuk menjalankan aplikasi PHP. Sistem dependency dan "
    "ekstensi PHP yang dibutuhkan diinstal di dalam image, termasuk pdo_mysql untuk koneksi ke MariaDB. "
    "Setelah source code disalin, Composer dijalankan untuk memasang dependency aplikasi. Container "
    "kemudian dijalankan melalui entrypoint script agar migrasi database dilakukan saat database sudah "
    "siap."
)
code_block(
    "FROM php:8.3-cli-alpine\n"
    "RUN apk add --no-cache git curl libpng-dev libxml2-dev zip unzip \\\n"
    "    sqlite-dev oniguruma-dev linux-headers mariadb-client $PHPIZE_DEPS\n"
    "RUN docker-php-ext-install pdo pdo_sqlite pdo_mysql mbstring exif pcntl bcmath gd xml\n"
    "COPY --from=composer:latest /usr/bin/composer /usr/bin/composer\n"
    "WORKDIR /var/www\n"
    "COPY . .\n"
    "RUN cp .env.example .env\n"
    "RUN composer install --optimize-autoloader --no-interaction\n"
    "RUN php artisan key:generate --force\n"
    "COPY docker/entrypoint.sh /var/www/docker/entrypoint.sh\n"
    "EXPOSE 8000\n"
    "CMD [\"sh\", \"/var/www/docker/entrypoint.sh\"]"
)

heading("4.2 Docker Compose", level=2)
para(
    "Docker Compose mengatur dua service utama, yaitu app dan db. Service app membangun image dari "
    "Dockerfile, memetakan port 8080 ke port internal 8000, dan mengakses database melalui hostname db. "
    "Service db menggunakan image MariaDB 11.4 dan menyimpan datanya pada volume db_data. Kredensial "
    "kedua service diambil dari environment variable (${DB_DATABASE}, ${DB_USERNAME}, ${DB_PASSWORD}, "
    "${DB_ROOT_PASSWORD}) yang bersumber dari file .env, bukan dituliskan langsung di dalam file yang "
    "dikirim ke repository."
)
code_block(
    "  db:\n"
    "    image: mariadb:11.4\n"
    "    environment:\n"
    "      MYSQL_ROOT_PASSWORD: ${DB_ROOT_PASSWORD}\n"
    "      MYSQL_DATABASE: ${DB_DATABASE}\n"
    "      MYSQL_USER: ${DB_USERNAME}\n"
    "      MYSQL_PASSWORD: ${DB_PASSWORD}\n"
    "    volumes:\n"
    "      - db_data:/var/lib/mysql\n"
    "    healthcheck:\n"
    "      test: [\"CMD-SHELL\", \"mariadb ... -u$$MYSQL_USER -p$$MYSQL_PASSWORD ...\"]\n\n"
    "  app:\n"
    "    build: .\n"
    "    ports: [\"8080:8000\"]\n"
    "    environment:\n"
    "      DB_HOST: db\n"
    "      DB_DATABASE: ${DB_DATABASE}\n"
    "      DB_USERNAME: ${DB_USERNAME}\n"
    "      DB_PASSWORD: ${DB_PASSWORD}\n"
    "    depends_on:\n"
    "      db:\n"
    "        condition: service_healthy\n"
    "    restart: unless-stopped\n\n"
    "volumes:\n"
    "  db_data:"
)

heading("4.3 Network dan Komunikasi Antar Service", level=2)
para(
    "Docker network todo-network digunakan agar service aplikasi dan database dapat saling "
    "berkomunikasi di dalam jaringan internal Docker. Aplikasi tidak perlu mengetahui IP database "
    "secara statis, cukup menggunakan nama service db sebagai hostname (DB_HOST=db). Cara ini lebih "
    "stabil dan sesuai dengan pola deployment container yang umum digunakan. Network internal juga "
    "memisahkan komunikasi aplikasi dari akses luar yang tidak diperlukan, karena port database 3306 "
    "tidak dipetakan ke host."
)

heading("4.4 Persistent Volume", level=2)
para(
    "Persistent volume diterapkan pada service database melalui volume db_data yang dipetakan ke "
    "/var/lib/mysql di dalam container. Volume ini menyimpan data MariaDB di luar siklus hidup "
    "container, sehingga data tidak hilang ketika container dihentikan atau dibuat ulang dengan "
    "docker compose down lalu docker compose up -d. Penggunaan volume merupakan komponen penting "
    "dalam cloud computing karena menjaga ketersediaan data dan mengurangi risiko kehilangan informasi "
    "akibat perubahan container. Tanpa volume, database akan bersifat sementara dan tidak cocok untuk "
    "skenario penyimpanan data produksi. Bukti pengujian persistent volume disajikan pada Bab 6.4."
)

heading("4.5 Health Check dan Ketahanan Layanan", level=2)
para(
    "Health check ditambahkan pada service database dan aplikasi. Pada database, health check "
    "menjalankan query SELECT 1 menggunakan kredensial dari environment variable container ($$MYSQL_USER "
    "/ $$MYSQL_PASSWORD) untuk memastikan MariaDB sudah siap menerima koneksi sebelum aplikasi dijalankan "
    "(depends_on dengan condition: service_healthy). Pada aplikasi, endpoint /api/health digunakan untuk "
    "memastikan aplikasi dan koneksi database tetap sehat, dan akan mengembalikan status 503 apabila "
    "koneksi database gagal. Restart policy unless-stopped juga diterapkan pada kedua service agar "
    "layanan dapat kembali aktif jika mengalami gangguan tak terduga. Kombinasi ini membuat sistem lebih "
    "tangguh dan mudah dipantau saat demo maupun saat dijalankan berulang kali. Bukti simulasi gangguan "
    "layanan disajikan pada Bab 6.5."
)


# ============================================================
# BAB 5 IMPLEMENTASI CI/CD
# ============================================================
heading("BAB 5 IMPLEMENTASI CI/CD", level=1)

heading("5.1 Tujuan Pipeline", level=2)
para(
    "Pipeline CI/CD digunakan untuk memastikan bahwa setiap perubahan kode melewati proses validasi "
    "otomatis sebelum dianggap siap digunakan. Pada proyek ini, GitHub Actions digunakan untuk "
    "menjalankan testing, build Docker image, dan verifikasi health check. Tujuan utama pipeline adalah "
    "menjaga kualitas kode, mengurangi kesalahan manual, dan membuktikan bahwa image yang dibangun dapat "
    "dijalankan dengan benar."
)

heading("5.2 Alur Workflow", level=2)
para(
    "Workflow .github/workflows/ci.yml dipicu pada setiap push dan pull request ke branch main. "
    "Workflow dimulai dengan checkout source code, kemudian setup PHP 8.3 dan instalasi dependency "
    "Composer (dengan cache). Setelah itu, file environment disiapkan (cp .env.example .env), "
    "application key dibuat, dan database SQLite disiapkan khusus untuk pengujian di CI sehingga tidak "
    "menyentuh data produksi. Berikutnya, automated testing dijalankan menggunakan PHPUnit. Jika test "
    "berhasil, workflow melanjutkan ke proses build Docker image, menyalakan stack Compose penuh "
    "(app + db), lalu memverifikasi health check dan endpoint yang terhubung ke database sebelum stack "
    "dimatikan kembali."
)
code_block(
    "jobs:\n"
    "  test-and-build:\n"
    "    runs-on: ubuntu-latest\n"
    "    steps:\n"
    "      - uses: actions/checkout@v4\n"
    "      - uses: shivammathur/setup-php@v2\n"
    "        with: { php-version: '8.3' }\n"
    "      - run: composer install --no-progress --prefer-dist\n"
    "      - run: cp .env.example .env && php artisan key:generate --force\n"
    "      - run: touch database/database.sqlite\n"
    "      - run: php artisan migrate --force      # DB_CONNECTION=sqlite (testing)\n"
    "      - run: php artisan test\n"
    "      - run: docker build -t todo-app:latest .\n"
    "      - run: docker compose up -d --build\n"
    "      - run: curl -fsS http://localhost:8080/api/health\n"
    "      - run: curl -fsS http://localhost:8080/api/todos\n"
    "      - run: docker compose down -v"
)

heading("5.3 Automated Testing sebagai Quality Gate", level=2)
para(
    "Automated testing berperan sebagai quality gate sebelum Docker image dinyatakan layak dibangun "
    "dan dijalankan. Test yang tersedia mencakup list todo, create todo, show todo, update todo, delete "
    "todo, dan validasi title wajib diisi (total 6 test khusus TodoApiTest, ditambah 2 test bawaan "
    "Laravel, sehingga 8 test berjalan). Dengan test ini, perubahan pada controller, route, atau "
    "validasi dapat terdeteksi lebih cepat sebelum masuk ke tahap build image maupun deployment. Test "
    "sengaja dijalankan di atas database SQLite yang terisolasi (bukan database produksi MariaDB), agar "
    "proses pengujian tidak berisiko mengubah atau menghapus data nyata milik pengguna."
)

heading("5.4 Simulasi Pipeline Gagal dan Berhasil", level=2)
para(
    "Sebagai bukti penerapan CI/CD, pipeline gagal pernah disimulasikan secara terkontrol melalui "
    "perubahan pada test (commit “Simulasi pipeline gagal”), sehingga GitHub Actions menampilkan "
    "status merah dan log kesalahan yang jelas. Setelah itu, kesalahan diperbaiki (commit “Memperbaiki "
    "test pipeline”) dan pipeline dijalankan kembali hingga berhasil dengan status hijau. Bukti pipeline "
    "gagal dan berhasil ini penting karena menunjukkan bahwa workflow benar-benar dipakai untuk "
    "memverifikasi kualitas codebase, bukan hanya sekadar file konfigurasi yang tidak pernah diuji."
)


# ============================================================
# BAB 6 PENGUJIAN
# ============================================================
heading("BAB 6 PENGUJIAN", level=1)
para(
    "Catatan: pada saat bukti berikut diambil, port host sementara dialihkan dari 8080 ke 8081 "
    "karena port 8080 pada komputer pengembang sedang dipakai oleh proyek Docker lain yang tidak "
    "berkaitan. Pengalihan ini hanya bersifat sementara untuk pengambilan bukti dan tidak mengubah "
    "konfigurasi docker-compose.yml yang dikumpulkan (tetap menggunakan port 8080:8000).",
    italic=True, size=10,
)

heading("6.1 Pengujian Fungsional", level=2)
para(
    "Pengujian fungsional dilakukan untuk memastikan aplikasi dapat menjalankan fitur utama sesuai "
    "kebutuhan. Hasil pengujian menunjukkan bahwa aplikasi mampu menampilkan daftar todo, menambahkan "
    "todo baru, memperbarui todo, menghapus todo, dan menolak input yang tidak valid. Endpoint "
    "/api/health juga berhasil merespons status sehat saat aplikasi dan database berjalan normal."
)
code_block(read_evidence("02_stack_running_and_health.txt"))

heading("6.2 Pengujian Automated Test", level=2)
para(
    "Automated testing dijalankan dengan perintah php artisan test di dalam container sekali pakai "
    "yang terpisah dari database produksi (menggunakan SQLite khusus testing, meniru langkah yang sama "
    "seperti pada pipeline CI). Seluruh 8 test berhasil lulus dengan 18 assertion. Hasil ini menunjukkan "
    "bahwa logika aplikasi sudah memiliki perlindungan dasar terhadap regresi tanpa membahayakan data "
    "produksi yang sedang berjalan."
)
code_block(read_evidence("03_php_artisan_test.txt"))

heading("6.3 Pengujian Container dan Compose", level=2)
para(
    "Pengujian container dilakukan dengan menjalankan docker compose up -d --build dan memeriksa "
    "versi serta status service menggunakan docker --version, docker compose version, dan "
    "docker compose ps. Hasil pengujian menunjukkan bahwa container app dan db dapat berjalan dengan "
    "status healthy. Aplikasi dapat diakses melalui port yang dipetakan, dan endpoint API dapat "
    "merespons dengan benar."
)
code_block(read_evidence("01_docker_version_and_ps.txt"))
code_block(read_evidence("09_docker_images.txt"))

heading("6.4 Pengujian Persistent Volume", level=2)
para(
    "Pengujian persistent volume dilakukan dengan menambahkan data baru ke database melalui endpoint "
    "POST /api/todos, kemudian menjalankan docker compose down (tanpa flag -v, sehingga volume tidak "
    "ikut dihapus), lalu menjalankan docker compose up -d kembali. Setelah container hidup kembali, data "
    "yang baru ditambahkan tetap tersedia karena disimpan pada volume db_data. Hasil ini membuktikan "
    "bahwa volume berfungsi sebagai mekanisme persistensi data, sehingga database tidak kehilangan isi "
    "saat lifecycle container berubah."
)
para("Data sebelum docker compose down:", bold=True)
code_block(read_evidence("04_persistence_before_down.txt"))
para("Proses docker compose down (volume db_data tetap ada):", bold=True)
code_block(read_evidence("05_compose_down.txt"))
para("Data setelah docker compose up -d kembali (identik dengan sebelum down):", bold=True)
code_block(read_evidence("06_persistence_after_up.txt"))

heading("6.5 Pengujian Ketahanan Layanan (Health Check)", level=2)
para(
    "Ketahanan layanan diuji dengan menghentikan container database secara paksa (docker stop todo-db) "
    "untuk mensimulasikan gangguan layanan. Endpoint /api/health pada aplikasi terbukti mendeteksi "
    "kegagalan koneksi database dan mengembalikan status unhealthy dengan kode HTTP 503, alih-alih "
    "menampilkan error yang tidak terkontrol. Setelah database dinyalakan kembali (docker start "
    "todo-db), endpoint /api/health kembali melaporkan status healthy dengan kode HTTP 200 dalam waktu "
    "singkat, menunjukkan bahwa layanan mampu pulih tanpa perlu membangun ulang container maupun "
    "kehilangan data."
)
para("Kondisi saat database dihentikan:", bold=True)
code_block(read_evidence("07_resilience_db_down.txt"))
para("Kondisi setelah database dinyalakan kembali:", bold=True)
code_block(read_evidence("08_resilience_recovery.txt"))


# ============================================================
# BAB 7 KESIMPULAN
# ============================================================
heading("BAB 7 KESIMPULAN", level=1)

heading("7.1 Kesimpulan Umum", level=2)
para(
    "Proyek Todo List ini berhasil mengimplementasikan aplikasi sederhana dalam arsitektur "
    "multi-container menggunakan Docker, Docker Compose, MariaDB, health check, persistent volume, "
    "automated testing, dan GitHub Actions. Walaupun aplikasi yang dikembangkan tidak kompleks, "
    "struktur implementasinya sudah mencerminkan praktik dasar cloud computing yang baik, terutama pada "
    "aspek containerization, orchestration, dan automation. Dengan demikian, proyek ini memenuhi fokus "
    "utama UAS, yaitu penerapan konsep cloud computing secara utuh dan dapat dibuktikan."
)

heading("7.2 Kendala dan Solusi", level=2)
para(
    "Kendala utama pada awal pengembangan adalah penggunaan database SQLite tunggal yang belum "
    "memenuhi syarat multi-container. Kendala tersebut diselesaikan dengan memindahkan runtime database "
    "ke MariaDB dalam container terpisah dan menambahkan volume untuk persistensi data. Kendala lain "
    "adalah memastikan aplikasi tidak start sebelum database siap, sehingga ditambahkan health check dan "
    "entrypoint script yang menunggu migrasi berhasil. Ditemukan pula bahwa menjalankan automated test "
    "langsung di dalam container produksi yang sedang berjalan berisiko mencampur data uji dengan data "
    "nyata; solusinya adalah menjalankan test pada container sekali pakai (docker compose run) dengan "
    "koneksi SQLite terpisah, sama seperti yang dilakukan pipeline CI. Kredensial database yang semula "
    "dituliskan langsung di docker-compose.yml juga dipindahkan menjadi environment variable yang "
    "bersumber dari file .env agar tidak tersimpan sebagai plaintext di repository."
)

heading("7.3 Rencana Pengembangan", level=2)
para(
    "Jika proyek ini dikembangkan lebih lanjut, service tambahan seperti Redis, Nginx, atau "
    "phpMyAdmin dapat ditambahkan untuk memperkaya arsitektur. Selain itu, backup database, logging "
    "terpusat, dan publikasi image ke Docker Hub atau GitHub Container Registry juga dapat diterapkan "
    "sebagai nilai tambah. Dalam skenario cloud, aplikasi ini juga dapat dipindahkan ke VPS atau layanan "
    "cloud dengan perubahan konfigurasi yang relatif kecil karena arsitekturnya sudah berbasis container; "
    "langkah yang diperlukan umumnya berupa penyesuaian DNS/reverse proxy, penggantian volume lokal "
    "dengan managed storage, dan penambahan secret management pada level platform cloud yang dituju."
)

# ============================================================
# LAMPIRAN
# ============================================================
heading("LAMPIRAN", level=1)
para("Bukti eksekusi lengkap (raw output) tersedia pada folder evidence/ di repository:", bold=False)
lampiran_items = [
    "evidence/01_docker_version_and_ps.txt - docker --version, docker compose version, docker compose ps",
    "evidence/02_stack_running_and_health.txt - stack berjalan healthy, /api/health, /api/todos",
    "evidence/03_php_artisan_test.txt - hasil php artisan test (8 test, isolated SQLite)",
    "evidence/04_persistence_before_down.txt - data sebelum docker compose down",
    "evidence/05_compose_down.txt - docker compose down, volume db_data tetap ada",
    "evidence/06_persistence_after_up.txt - data setelah docker compose up -d kembali",
    "evidence/07_resilience_db_down.txt - simulasi database mati, /api/health -> 503 unhealthy",
    "evidence/08_resilience_recovery.txt - database pulih, /api/health -> 200 healthy",
    "evidence/09_docker_images.txt - docker images (image todo-app-app berhasil dibangun)",
    "docs/architecture-diagram.svg / .png - diagram arsitektur (sumber vektor dan raster)",
]
for item in lampiran_items:
    para(item, size=10)

doc.add_paragraph()
heading("Format Pengumpulan", level=2)
fmt_items = [
    "Nama/NIM/Kelas: Purna Siswantomo / C2C023160 / [Kelas Anda]",
    "Nama/NIM/Kelas: Erifa Dwi Astuti / C2C023161 / [Kelas Anda]",
    "Nama aplikasi: Todo List Application",
    "Link repository GitHub: https://github.com/Purna-Siswantomo/Todo-docker",
    "Link pipeline gagal: https://github.com/Purna-Siswantomo/Todo-docker/actions/runs/28741426116",
    "Link pipeline berhasil: https://github.com/Purna-Siswantomo/Todo-docker/actions/runs/28741450472",
    "Link video demonstrasi: [diisi setelah video direkam dan diunggah]",
    "Link image registry (opsional): [belum tersedia]",
    "Link aplikasi online (opsional): [belum tersedia]",
]
for item in fmt_items:
    para(item, size=11)

OUT = REPO / "C2C023160_PurnaSiswantomo_UAS_CloudComputing.docx"
doc.save(str(OUT))
print("Saved:", OUT)
